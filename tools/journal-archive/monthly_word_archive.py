import argparse, json, re
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

LABELS = {
    "date":"Дата", "createdAt":"Создано", "updatedAt":"Обновлено", "startedAt":"Начало",
    "endedAt":"Окончание", "resolvedAt":"Устранено", "confirmedAt":"Подтверждено",
    "name":"Наименование", "area":"Участок", "equipmentName":"Оборудование", "equipmentId":"ID оборудования",
    "node":"Узел", "nodeName":"Узел", "shift":"Смена", "status":"Статус", "comment":"Комментарий",
    "description":"Описание", "resolution":"Устранение", "resolutionComment":"Комментарий об устранении",
    "createdBy":"Кто внёс", "createdByName":"Кто внёс", "resolvedBy":"Кто устранил",
    "resolvedByName":"Кто устранил", "confirmedBy":"Кто подтвердил", "confirmedByName":"Кто подтвердил",
    "operatorName":"Оператор", "employeeNumber":"Табельный №", "role":"Должность/роль",
    "installedPart":"Установленная запчасть", "partName":"Запчасть", "quantity":"Количество",
    "result":"Результат", "remarks":"Замечания", "type":"Вид", "mode":"Вид осмотра"
}

CATEGORIES = [
    ("QR-обходы", "qrWalkJournal"),
    ("Предупреждения и устранения", "checks"),
    ("Заявки и ремонты", "requests"),
    ("ППР", "pprSheets"),
    ("Простои", "downtimes"),
    ("Компрессорный журнал", "compressorJournal"),
    ("Журнал ШГРП и ГРП", "gasJournal"),
    ("Журнал сварочных работ", "weldingJournal"),
    ("Журнал токарных работ", "turningJournal"),
    ("Затраты и обслуживание", "serviceCosts"),
]

DATE_KEYS = ("date","at","createdAt","updatedAt","startedAt","endedAt","resolvedAt","confirmedAt","completedAt","inspectionAt")

HUMAN_LABELS = {
    "at":"Дата и время", "completedAt":"Работа завершена", "inspectionAt":"Время осмотра",
    "equipment":"Оборудование", "capacity":"Грузоподъёмность", "location":"Место установки",
    "authorName":"Работник", "authorRole":"Должность", "byName":"Работник", "byRole":"Должность",
    "shiftDate":"Дата", "shiftLabel":"Смена", "inspectionType":"Вид осмотра",
    "decision":"Итог осмотра", "defects":"Замечания", "points":"Результаты проверки",
    "checkedBy":"Проверил", "fixedByName":"Зафиксировал", "fixedAt":"Дата фиксации",
    "updatedByName":"Последнее изменение — Ф.И.О.", "updatedByRole":"Должность",
    "operatingState":"Рабочее состояние", "airPressure":"Давление воздуха",
    "airTemp":"Температура воздуха", "oilPressureTemp":"Давление и температура масла",
    "leakGrounding":"Утечки и заземление", "blowTime":"Время продувки",
    "equipmentStatus":"Состояние оборудования", "gasSmell":"Запах газа",
    "inletMpa":"Давление на входе, МПа", "outletMpa":"Давление на выходе, МПа",
    "pressureDeltaMpa":"Перепад давления, МПа", "protectionZone":"Охранная зона",
    "pskTrigger":"Срабатывание ПСК", "maintenance":"Техническое обслуживание",
    "actions":"Выполненные действия", "route":"Маршрут", "section":"Участок",
    "time":"Время", "shiftTime":"Время смены", "entryStatus":"Состояние записи",
    "closeComment":"Комментарий о завершении", "closedByName":"Кто завершил",
    "closedByRole":"Должность завершившего", "createdByName":"Автор",
    "createdByPosition":"Должность автора", "description":"Описание работы",
    "drawingNumber":"Номер чертежа", "material":"Материал", "participants":"Исполнители",
    "workComment":"Результат и комментарий", "welderName":"Сварщик",
    "welderPosition":"Должность сварщика", "welderCertificate":"Удостоверение сварщика",
    "welderStamp":"Клеймо сварщика", "jointPosition":"Положение шва",
    "consumables":"Расходные материалы", "turnerName":"Токарь", "machine":"Станок",
    "operations":"Выполненные операции", "measurements":"Контрольные размеры",
    "quantity":"Количество", "madeQty":"Изготовлено", "goodQty":"Годных",
    "rejectQty":"Брак", "rejectReason":"Причина брака", "blankSize":"Размер заготовки",
    "dueDate":"Срок выполнения", "plannedAt":"Плановая дата", "plannedByName":"Запланировал",
    "approvedAt":"Дата утверждения", "approvedByName":"Утвердил", "rows":"Перечень работ",
    "group":"Группа обхода", "source":"Источник записи"
}
LABELS.update(HUMAN_LABELS)

HIDDEN_FIELDS = {
    "id","key","equipmentId","sourceEquipmentId","sourceRecordKey","authorKey",
    "authorEmployeeId","createdById","resolvedById","confirmedById","closedByKey","welderId",
    "turnerId","inspectorKey","inspectorKeys","engineerKeys","shiftKey","archivedNodeIndex",
    "nodeIndex","passwordHash","photos","photo","image","requestPhoto","resultPhoto",
    "autofillInitialized","autofillMode","autofilledFor","plannedAutomatically","deleted",
    "deletedAt","duplicateOf"
}

VALUE_LABELS = {
    "allowed":"К работе допущено", "blocked":"Эксплуатация запрещена",
    "day":"Ежесменный осмотр", "monthly":"Ежемесячный осмотр электромехаником",
    "ok":"Исправно", "remark":"Есть замечание", "resolved":"Устранено",
    "pending":"Ожидает подтверждения", "completed":"Выполнено", "open":"Открыто"
}

def safe_name(value):
    text = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', '_', str(value or '')).strip(' ._')
    return (text[:100] or 'Журнал')

def nested(data, path):
    cur=data
    for key in path.split('.'):
        if not isinstance(cur, dict): return {}
        cur=cur.get(key,{})
    return cur

def records(value):
    if isinstance(value,list): return [(str(i+1),x) for i,x in enumerate(value)]
    if isinstance(value,dict): return list(value.items())
    return []

def month_match(key, item, month):
    blob = f"{key} {json.dumps(item, ensure_ascii=False, default=str)}"
    return month in blob

def compact(value, depth=0):
    if value is None or value == "": return "—"
    if isinstance(value,bool): return "Да" if value else "Нет"
    if isinstance(value,(int,float)): return str(value)
    if isinstance(value,list):
        return "; ".join(compact(x,depth+1) for x in value[:30]) or "—"
    if isinstance(value,dict):
        if depth > 1: return json.dumps(value,ensure_ascii=False,default=str)
        return "; ".join(f"{LABELS.get(k,k)}: {compact(v,depth+1)}" for k,v in value.items() if v not in (None,"",[],{})) or "—"
    text = str(value)
    if re.match(r'^\d{4}-\d{2}-\d{2}T', text):
        try:
            return datetime.fromisoformat(text.replace('Z', '+00:00')).strftime('%d.%m.%Y %H:%M')
        except ValueError:
            pass
    if re.match(r'^\d{4}-\d{2}-\d{2}$', text):
        try:
            return datetime.strptime(text, '%Y-%m-%d').strftime('%d.%m.%Y')
        except ValueError:
            pass
    return VALUE_LABELS.get(text, text)

def set_repeat_header(row):
    trPr=row._tr.get_or_add_trPr(); el=OxmlElement('w:tblHeader'); el.set(qn('w:val'),'true'); trPr.append(el)

def add_record(doc, number, key, item):
    p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(3); p.paragraph_format.keep_with_next=True
    r=p.add_run(f"Запись {number}"); r.bold=True; r.font.size=Pt(11)
    if not isinstance(item,dict): item={"Значение":item}
    preferred=[k for k in DATE_KEYS if item.get(k)]
    keys=preferred+[k for k in item if k not in preferred and k not in HIDDEN_FIELDS and k in LABELS]
    table=doc.add_table(rows=0,cols=2); table.style='Table Grid'; table.autofit=False
    for field in keys:
        val=item.get(field)
        if val in (None,"",[],{}): continue
        row=table.add_row(); row.cells[0].width=Cm(4.5); row.cells[1].width=Cm(12)
        row.cells[0].text="Ключ записи" if field=="_record_key" else LABELS.get(field,field)
        row.cells[1].text=compact(val)
        for rr in row.cells[0].paragraphs[0].runs: rr.bold=True
        for cell in row.cells:
            for pp in cell.paragraphs:
                pp.paragraph_format.space_after=Pt(0)
                for rr in pp.runs: rr.font.name='Arial'; rr.font.size=Pt(8.5)

def make_doc(title, month, items, out_file, company):
    doc=Document(); sec=doc.sections[0]
    sec.top_margin=Cm(1.5); sec.bottom_margin=Cm(1.5); sec.left_margin=Cm(1.7); sec.right_margin=Cm(1.7)
    style=doc.styles['Normal']; style.font.name='Arial'; style.font.size=Pt(10); style.paragraph_format.space_after=Pt(4)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(title.upper()); r.bold=True; r.font.name='Arial'; r.font.size=Pt(15)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{company} · месяц {month} · записей: {len(items)}")
    for i,(key,item) in enumerate(items,1): add_record(doc,i,key,item)
    if not items: doc.add_paragraph("За выбранный месяц записей нет.")
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Сформировано автоматически: "+datetime.now().strftime("%d.%m.%Y %H:%M"))
    doc.save(out_file)

def main():
    ap=argparse.ArgumentParser(); ap.add_argument('--input',required=True); ap.add_argument('--month',required=True); ap.add_argument('--output',required=True)
    a=ap.parse_args(); data=json.loads(Path(a.input).read_text(encoding='utf-8-sig'))
    company=((data.get('adminConfig') or {}).get('companyName') or 'Организация')
    root=Path(a.output); root.mkdir(parents=True,exist_ok=True)
    manifest=[]
    for title,path in CATEGORIES:
        selected=[(k,v) for k,v in records(nested(data,path)) if month_match(k,v,a.month)]
        if not selected: continue
        folder=root/safe_name(title); folder.mkdir(exist_ok=True)
        out=folder/f"{safe_name(title)}_{a.month}.docx"
        make_doc(title,a.month,selected,out,company); manifest.append({"journal":title,"records":len(selected),"file":str(out)})
    (root/'Манифест.json').write_text(json.dumps({"month":a.month,"createdAt":datetime.now().isoformat(),"files":manifest},ensure_ascii=False,indent=2),encoding='utf-8')
    (root/'READY.txt').write_text(f"Месяц: {a.month}\nФайлов Word: {len(manifest)}\n",encoding='utf-8')
    print(json.dumps({"month":a.month,"files":len(manifest),"output":str(root)},ensure_ascii=False))

if __name__=='__main__': main()
